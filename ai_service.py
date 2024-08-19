from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
from langchain_community.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_community.document_loaders import TextLoader
import shutil
import os
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from paddleocr import PaddleOCR
import fitz
from PIL import Image
from tqdm import tqdm
import random
from llama_parse import LlamaParse
load_dotenv()

llm = ChatOpenAI(model='gpt-4o', temperature=0.0)

system_prompt = """# OBJECTIVE:
As a Medical & Legal Assistant of MEDirect, your role encompasses summarizing medical documents and providing clear and comprehensive information necessary for legal and medical evaluation. The medical documents you will process include Letter Of Instruction (LOI), Independent Medical Examination (IME) report and other medical & legal documents.
As a Medical Assistant, your responsibility is to extract 8 or 12 formatted paragraphs of high level summarization of medical document provided without any emotional or subjective language.

# TONE:
Provide the summarization in formal, clinical, factual, precise and neutral tone.

# IMPORTANT NOTES:
1. Format the response in several paragraphs.
2. Do not add any title or heading, just the paragraphs of summarization directly.
3. Answer directly without any introductions.

{context}
"""

user_prompt_for_background_history = """
Extract Background Information from the following letter of instruction including all important information like key personal details, employment history, medical history and treatment and legal and compensation details are still included in the summarization focusing on the events, diagnoses, treatments and current condition.
"""

user_prompt_for_ime = """
Extract the name of the doctor with his or her position who conducted assessment or examination from the following medical document and summarize the content of the document in super details.
The summarization must contain the dates mentioned in the document including all the details and key info.
Make sure that the FIRST line of the response should be the name of the doctor with its position and from the SECOND line should start the summarization.
"""

user_prompt_for_questions = """
Extract specific questions with numbering for examination from the following letter of instruction.
"""

prompt_template = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "{input}")
    ])

BACKGROUND_INFORMATION_HEADING = "BACKGROUND INFORMATION"
IME_SUMMARY_HEADING = "DOCUMENT REVIEW"
SPECIFIC_QUESTIONS_HEADING = "Attachment 1: SPECIFIC QUESTIONS FOR EXAMINATION"

HEADINGS_TO_INSERT_LOI = ["BACKGROUND INFORMATION", "Attachment 1: SPECIFIC QUESTIONS FOR EXAMINATION"]
HEADINGS_TO_INSERT_IME = ["DOCUMENT REVIEW"]

TEMP_VECTOR_DB_FOLDER = "temp_index"
TEMP_IMAGE_FOLDER = 'temp_images'
TEMP_TEXT_FOLDER = 'temp_text'

def create_vector_store_from_txt(text_path):
    if not os.path.exists(TEMP_VECTOR_DB_FOLDER):
        os.makedirs(TEMP_VECTOR_DB_FOLDER)
    
    file_name = os.path.basename(text_path).split('.')[0]
    vector_store_path = os.path.join(TEMP_VECTOR_DB_FOLDER, file_name)
    loader = TextLoader(text_path)
    pages = loader.load_and_split()
    if not pages:
        raise ValueError("No valid text data to create vector store.")
    vector_store = FAISS.from_documents(pages, OpenAIEmbeddings())
    vector_store.save_local(vector_store_path)
    print('Vector Store Created..')
    
    return vector_store_path

def retrieve_info(vector_db_path, user_prompt):
    vector_db = FAISS.load_local(vector_db_path, OpenAIEmbeddings(), allow_dangerous_deserialization=True)
    retriever = vector_db.as_retriever()
    stuff_documents_chain = create_stuff_documents_chain(llm, prompt_template)
    rag_chain = create_retrieval_chain(retriever, stuff_documents_chain)
    response = rag_chain.invoke({
        "input": user_prompt
    })["answer"]
    
    return response

def generate_answer_loi(loi_vector_db, heading):
    if heading == "BACKGROUND INFORMATION":
        return retrieve_info(loi_vector_db, user_prompt_for_background_history)      
    elif heading == "Attachment 1: SPECIFIC QUESTIONS FOR EXAMINATION":
        return retrieve_info(loi_vector_db, user_prompt_for_questions)  

def generate_answer_ime(ime_vector_db, heading):
    if heading == "DOCUMENT REVIEW":
        return retrieve_info(ime_vector_db, user_prompt_for_ime)      

def insert_from_loi(loi_vector_db, report_base):
    for heading in HEADINGS_TO_INSERT_LOI:        
        for paragraph in report_base.paragraphs:
            if heading in paragraph.text:                
                insertion = generate_answer_loi(loi_vector_db, heading)                
                new_paragraph = OxmlElement("w:p")
                paragraph._p.addnext(new_paragraph)
                final_paragraph = Paragraph(new_paragraph, paragraph._parent)
                run = final_paragraph.add_run(f'\n{insertion}')
                run.font.size = Pt(12)
                run.font.name = 'Arial'            
                break

def insert_from_ime(ime_vector_db, report_base):
    for heading in HEADINGS_TO_INSERT_IME:        
        for paragraph in report_base.paragraphs:
            if heading in paragraph.text:                
                insertion = generate_answer_ime(ime_vector_db, heading)
                insertion_lines = insertion.split('\n')
                                
                new_paragraph = OxmlElement("w:p")
                paragraph._p.addnext(new_paragraph)
                final_paragraph = Paragraph(new_paragraph, paragraph._parent)
                
                # Add the first line in bold
                if insertion_lines:
                    first_line_run = final_paragraph.add_run(f'\n{insertion_lines[0]}')
                    first_line_run.bold = True
                    first_line_run.font.size = Pt(12)
                    first_line_run.font.name = 'Arial'
                
                # Add the remaining lines as normal text
                if len(insertion_lines) > 1:
                    for line in insertion_lines[1:]:
                        remaining_line_run = final_paragraph.add_run(f'\n{line}')
                        remaining_line_run.font.size = Pt(12)
                        remaining_line_run.font.name = 'Arial'            
                break
            
def generate_final_report(doc_loi, doc_ime_report, doc_report_template, report_folder):
    if os.path.exists(doc_report_template):
        report_base = Document(doc_report_template)
    
    if os.path.exists(doc_loi):
        loi_text, loi_text_path = get_text_by_llamaparse(doc_loi)
        loi_vector_db = create_vector_store_from_txt(loi_text_path)
        insert_from_loi(loi_vector_db, report_base)
        os.remove(doc_loi)    
    
    for doc_ime_item in doc_ime_report:
        if os.path.exists(doc_ime_item):        
            ime_text, ime_text_path= get_text_by_llamaparse(doc_ime_item)    
            ime_vector_db = create_vector_store_from_txt(ime_text_path)
            insert_from_ime(ime_vector_db, report_base)
            os.remove(doc_ime_item)
    
    shutil.rmtree(TEMP_VECTOR_DB_FOLDER)
    shutil.rmtree(TEMP_TEXT_FOLDER)  
        
    generated_report_path = os.path.join(report_folder, f"{os.path.basename(doc_report_template).split('.')[0]}_inserted.docx")
    report_base.save(generated_report_path)
    print("Final Report Generated!")
    return generated_report_path

def get_text_from_pdf(pdf_path):
    if not os.path.exists(TEMP_IMAGE_FOLDER):
        os.makedirs(TEMP_IMAGE_FOLDER)
    if not os.path.exists(TEMP_TEXT_FOLDER):
        os.makedirs(TEMP_TEXT_FOLDER)
    
    doc_opened = fitz.open(pdf_path)
    #iterate through each pages and convert them into a png file
    for page_num in tqdm(range(len(doc_opened)), desc='Converting PDF Pages to Images: '):
        cur_page = doc_opened.load_page(page_num)
        render_pixmap = cur_page.get_pixmap()

        img = Image.frombytes("RGB", [render_pixmap.width, render_pixmap.height], render_pixmap.samples)

        out = os.path.join(TEMP_IMAGE_FOLDER, f'page_{page_num + 1}.png')
        img.save(out)        

    ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
    
    image_list = os.listdir(TEMP_IMAGE_FOLDER)
    image_list.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))
    detected_text = []
    image_count = 0
    for image in tqdm(image_list, desc='Reading texts from Images: ', total=len(image_list)):
        image_path = os.path.join(TEMP_IMAGE_FOLDER, image)
        start_decode = ocr.ocr(image_path, cls=True) 
        image_count += 1

        bbox_list = []
        cnt_vertical = 0

        temp_text = []

        #compile all the texts for each image
        for line in start_decode:
            try:
                for bbox, (text, confidence) in line:
                    bbox_list.append(bbox)
                    #append the predicted text into the list
                    temp_text.append(text)
            except TypeError: #Nonetype
                #no words on current line
                #will produce None as output and will cause TypeError unless handled
                continue

        # Determine if this image is vertical or not
        cnt_check = int(len(bbox_list)*4/10) if len(bbox_list) > 30 else len(bbox_list)
        bboxes_to_check = random.choices(bbox_list, k=cnt_check)
        for bbox in bboxes_to_check:
            width = abs(bbox[1][0] - bbox[0][0])
            height = abs(bbox[2][1] - bbox[1][1])
            if height > width:
                cnt_vertical += 1
        
        try:
            if float(cnt_vertical/cnt_check) > 0.5:                
                image_ = Image.open(image_path)
                rotated_image_ = image_.rotate(90, expand=True)
                rotated_image_.save(image_path)                
                start_decode = ocr.ocr(image_path, cls=True)

                temp_text = []
                for line in start_decode:                   
                    try:
                        for bbox, (text, confidence) in line:
                            bbox_list.append(bbox)
                            #append the predicted text into the list
                            temp_text.append(text)
                    except TypeError: #Nonetype
                        #no words on current line
                        #will produce None as output and will cause TypeError unless handled
                        continue
        except ZeroDivisionError:
            pass
            #ignore this part, this means that the page is blank
        
        detected_text.extend(temp_text)        

    combined_text = ' '.join(detected_text)
    txt_path = os.path.join(TEMP_TEXT_FOLDER, f"{os.path.basename(pdf_path).split('.')[0]}.txt")
    with open(txt_path, 'w') as f:
        f.write(combined_text)
    shutil.rmtree(TEMP_IMAGE_FOLDER)
    
    return combined_text, txt_path

def get_text_by_llamaparse(pdf_path):
    if not os.path.exists(TEMP_TEXT_FOLDER):
        os.makedirs(TEMP_TEXT_FOLDER)
    parser = LlamaParse(
    api_key="llx-z54VQvmjOxCzLMUd5BFZApbo7OEVw1H4xNwLfmJD3WD5PSov",  # can also be set in your env as LLAMA_CLOUD_API_KEY
    result_type="text",  # "markdown" and "text" are available
    verbose=True,
    )
    
    parsed_result = ''
    documents = parser.load_data(pdf_path)
    
    for document in documents:
        parsed_result += document.text    
    txt_path = os.path.join(TEMP_TEXT_FOLDER, f"{os.path.basename(pdf_path).split('.')[0]}.txt")
    
    with open(txt_path, 'w') as f:
        f.write(parsed_result)
    
    return parsed_result, txt_path