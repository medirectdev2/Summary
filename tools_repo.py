from langchain_community.document_loaders import PyMuPDFLoader
#from langchain_community.document_loaders import PyPDFLoader
#from langchain.document_loaders.pdf import PyPDFDirectoryLoader
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv
load_dotenv()
from langchain_openai import ChatOpenAI
from langchain.tools import BaseTool
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader

from langchain.agents import create_openai_functions_agent, AgentExecutor
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import MessagesPlaceholder
from langchain_core.messages import SystemMessage

import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

#unique filename generator
import uuid


from paddleocr import PaddleOCR
import fitz
import os
from PIL import Image
from tqdm import tqdm
import random



llm = ChatOpenAI(model='gpt-4o', temperature=0.0)


def create_Vector(filepath):
    #this method is removed for now to test OCR implementation
    print('Creating Vector Store Database..')
    print(filepath)
    loader = PyMuPDFLoader(file_path=filepath)
    pages = loader.load_and_split()
    
    #generate unique temp folder
    
    

    vector_store = FAISS.from_documents(pages, OpenAIEmbeddings())
    vector_store.save_local('PDF_index')
    print('Vector Store Created..')



def test_PDF(question, parent_fold):
    print('running PDFTool')

    vector_store = FAISS.load_local(os.path.join(parent_fold, 'PDF_index'), OpenAIEmbeddings(), allow_dangerous_deserialization=True)

    retriever = vector_store.as_retriever()

    if question == 'BACKGROUND INFORMATION':
        input = '''
            1. Look for the Background Information and summarize the details inside it into a paragraph.
            2. Make sure that all important details are still included in the summary.      
            3. Be as detailed as possible when creating the summary.
            4. Do not add any title/header, just the paragraph summary directly. 
            5. Answer directly without any introductions.
        '''
    elif question == 'Attachment 1: SPECIFIC QUESTIONS FOR EXAMINATION':
        input = '''
            1. Copy all the questions listed from "Attachment 1:Specific questions for examination" including the spaces and tabs.
            2. Do not include the header 'Attachment 1:SPECIFIC QUESTIONS FOR EXAMINATION' when copying.
            3. Format the copied questions in number format or bullet format depending on how it's arranged.
        '''
    elif question == 'INVESTIGATIONS':
        input = '''
            1. Look for the Independent Medical Examination (IME) Report and summarize them.
            2. If there are multiple Independent Medical Examination (IME) Report, make sure you summarize each of them.
            3. Do not combine the summarization if there are multiple Independent Medical Examination (IME) Report.
            4. Answer directly without any introductions.
        '''
    elif question == 'HISTORY OF INJURY':
        input = '''
            1. Summarize the patient's injury history based on the information given.
            2. Answer directly in paragraph format
        '''
    else:
        input = f'''
            1. Look for the appropriate detail regarding this: Patient's {question}.
            3. Answer directly in paragraph format.
        '''

    prompt = ChatPromptTemplate.from_messages(
        [
            ("user", "{input}"),
            ("user", "Given the above conversation, generate a search query to look up in order to get information relevant to the conversation. If there are no relevant information, respond with None")
        ]
    )

    retriever_chain = create_history_aware_retriever(llm, retriever, prompt)


    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "Answer the user's questions based on the below context:\n\n{context}"),
            ("user", "{input}")
        ]
    )

    stuff_documents_chain = create_stuff_documents_chain(llm, prompt)

    finalchain = create_retrieval_chain(retriever_chain, stuff_documents_chain)

    getans = finalchain.invoke(
        {
            "input": input
        }
    )

    return(getans['answer'])



# tools_repo.py

from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

def start_generation(report_type, report_template, update_progress, filepath, parent_fold):
    print(report_type)

    # generated_uuid = uuid.uuid4()
    # parent_fold = f'Main_Folder_{generated_uuid}'
    image_temp = os.path.join(parent_fold, 'Image_Folder')
    if not os.path.exists(parent_fold):
        os.makedirs(parent_fold)

    if not os.path.exists(image_temp):
        os.makedirs(image_temp)

    open_doc = fitz.open(filepath)

    #debug attachment
    #if not os.path.exists(parent_fold)

    
    #iterate through each pages and convert them into a png file
    for page_num in tqdm(range(len(open_doc)), desc='Converting PDF Pages to Images: '):
        cur_page = open_doc.load_page(page_num)
        render_pixmap = cur_page.get_pixmap()

        img = Image.frombytes("RGB", [render_pixmap.width, render_pixmap.height], render_pixmap.samples)

        out = os.path.join(image_temp, f'page_{page_num + 1}.png')
        img.save(out)
        update_progress(f'Converting PDF to Image - ({page_num+1} of {len(open_doc)})')

    ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)

    
    image_list = os.listdir(image_temp)
    image_list.sort(key=lambda x: int(x.split('_')[1].split('.')[0]))
    detected_text = []
    image_count = 0
    for image in tqdm(image_list, desc='Reading texts from Images: ', total=len(image_list)):
        image_path = os.path.join(image_temp, image)
        start_decode = ocr.ocr(image_path, cls=True) 
        image_count += 1

        bbox_list = []
        cnt_vertical = 0

        temp_text = []

        #compile all the texts for each image
        for line in start_decode:
            try:
                for bbox, (text, confidence) in line:
                    bbox_list.append(bbox)
                    #append the predicted text into the list
                    temp_text.append(text)
            except TypeError: #Nonetype
                #no words on current line
                #will produce None as output and will cause TypeError unless handled
                continue

        # Determine if this image is vertical or not
        cnt_check = int(len(bbox_list)*4/10) if len(bbox_list) > 30 else len(bbox_list)
        bboxes_to_check = random.choices(bbox_list, k=cnt_check)
        for bbox in bboxes_to_check:
            width = abs(bbox[1][0] - bbox[0][0])
            height = abs(bbox[2][1] - bbox[1][1])
            if height > width:
                cnt_vertical += 1
        
        try:
            if float(cnt_vertical/cnt_check) > 0.5:
                #print(f"{image} is vertical!")
                image_ = Image.open(image_path)
                rotated_image_ = image_.rotate(90, expand=True)
                rotated_image_.save(image_path)
                #print("Rotated!")
                start_decode = ocr.ocr(image_path, cls=True)

                temp_text = []
                for line in start_decode:                   
                    try:
                        for bbox, (text, confidence) in line:
                            bbox_list.append(bbox)
                            #append the predicted text into the list
                            temp_text.append(text)
                    except TypeError: #Nonetype
                        #no words on current line
                        #will produce None as output and will cause TypeError unless handled
                        continue
        except ZeroDivisionError:
            pass
            #ignore this part, this means that the page is blank
        
        detected_text.extend(temp_text)

        update_progress(f'Gathering data from PDF - {image_count} of {len(image_list)}')

    combine_text = ' '.join(detected_text)
    with open(os.path.join(parent_fold, 'Generated_Text.txt'), 'w') as f:
        f.write(combine_text)

    loader = TextLoader(os.path.join(parent_fold, 'Generated_Text.txt'))
    pages = loader.load_and_split()

    print(len(pages))
    

    vector_store = FAISS.from_documents(pages, OpenAIEmbeddings())
    vector_store.save_local(os.path.join(parent_fold, 'PDF_index'))
    print('Vector Store Created..')


    # Open base doc file
    open_existing = Document(report_template)
    
    if report_type == 'Pain Assessment':
        details = ['BACKGROUND INFORMATION', 'HISTORY OF INJURY', 'SUBSEQUENT WORKER HISTORY',
                   'SOCIAL CIRCUMSTANCES', 'PAST HISTORY OF INJURY', 'MEDICAL HISTORY',
                   'JOB HISTORY', 'PRESENT COMPLAINTS', 'TREATMENT', 'EXAMINATION',
                   'INVESTIGATIONS', 'Attachment 1: SPECIFIC QUESTIONS FOR EXAMINATION']
        #details = ['BACKGROUND INFORMATION']
    elif report_type == 'Suplementary Report':
        details = ['BACKGROUND INFORMATION']
    elif report_type == 'Report Template - 130 Week' or report_type == 'Report Template - 130 Week - Woolworths Questions':
        details = ['BACKGROUND INFORMATION', 'DOCUMENT REVIEW', 'HISTORY OF INJURY', 'SUBSEQUENT WORKER HISTORY',
                   'SOCIAL CIRCUMSTANCES', 'PAST HISTORY OF INJURY', 'MEDICAL HISTORY',
                   'JOB HISTORY', 'PRESENT COMPLAINTS', 'Activities of daily living', 'TREATMENT', 'EXAMINATION',
                   'INVESTIGATIONS']
    elif report_type == 'Impairment Report':
        details = ['Description of accident / injury', 'HIstory since accident / injury including contribution of employment to current state.',
                   'Current symptoms', 'Medication (pre- and post-accident)', 'Past history',
                   'Lifestyle of worker including activities outside of work', 'Clinical examination (objective and observed versus formal examination range of movements)',
                   'Investigations / special tests', 'Analysis of findings', '•	Current diagnosis of injuries / conditions',
                   '•	Does the worker continue to suffer from any medical condition relevant to the alleged injury?']
    elif report_type == 'Case Notes':
        details = ['BACKGROUND', 'HISTORY OF PRESENTING COMPLAINT', 'PAST MEDICAL HISTORY',
                   'CURRENT MEDICATIONS', 'CURRENT COMPLAINTS', 'CURRENT TREATMENT',
                   'EXAMINATION', 'INVESTIGATIONS', 'SPECIFIC QUESTIONS']

    # Add the necessary generated answers
    for data in details:
        # Update progress on the loading page
        update_progress(f'Processing: {data}')

        for paragraph in open_existing.paragraphs:
            if data in paragraph.text:
                # Retrieve the generated response
                hey = test_PDF(data, parent_fold)
                
                new_p = OxmlElement("w:p")
                paragraph._p.addnext(new_p)
                new_para = Paragraph(new_p, paragraph._parent)
                run = new_para.add_run(f'\n{hey}')
                
                # Add a new paragraph with the generated response immediately after the matching paragraph
                # Set the font size and font name for the run
                run.font.size = Pt(12)
                run.font.name = 'Arial'

                # Stop searching after the first match to avoid multiple insertions
                break

    now = datetime.now()
    date = now.strftime("%Y-%m-%d")  # Date in the format YYYY-MM-DD
    hour = now.strftime("%H")  # Hour in 24-hour format
    minute = now.strftime("%M")  # Minute
    output_path = os.path.join(parent_fold, 'processed', f'{report_type}-{date}_{hour}_{minute}.docx')
    open_existing.save(output_path)
    print('All Done!')
    
    return output_path
    #add shutil rmtree somewhere to autodelete the folder
    #do not keep the data uploaded by the user


